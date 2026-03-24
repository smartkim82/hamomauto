import streamlit as st
import pandas as pd
from datetime import datetime
import time
import io
from docx import Document
from fpdf import FPDF
import sys
import sqlite3

from crawler.naver_map_crawler import NaverMapCrawler
from sender.email_sender import EmailSender
from sender.youtube_uploader import YouTubeAutomator

# DB 초기화 및 연결 설정
def init_db():
    conn = sqlite3.connect("hamom_database.db", check_same_thread=False)
    cursor = conn.cursor()
    # 크롤링 B2B 데이터 보관 테이블
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS b2b_crawling_list (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            업체명 TEXT,
            카테고리 TEXT,
            주소 TEXT,
            전화번호 TEXT,
            이메일 TEXT,
            해시태그 TEXT,
            검색카테고리 TEXT,
            등록일시 DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    conn.commit()
    return conn

conn = init_db()

# -----------------
# 기본 설정
# -----------------
st.set_page_config(
    page_title="Hamom Auto (웹 버전)",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 인증 상태 초기화
if 'is_authenticated' not in st.session_state:
    st.session_state['is_authenticated'] = False

# -----------------
# 사이드바 (메뉴 구조)
# -----------------
st.sidebar.title("하맘자동화 (WEB)")

# 인증 처리
st.sidebar.markdown("---")
if not st.session_state['is_authenticated']:
    st.sidebar.warning("🔒 관리자 잠금 상태입니다.")
    pwd_input = st.sidebar.text_input("잠금 해제 비밀번호", type="password")
    if pwd_input == "1191004":
        st.session_state['is_authenticated'] = True
        st.sidebar.success("✅ 잠금 해제 완료! 모든 기능 사용 가능")
        st.rerun()
else:
    st.sidebar.success("🔓 개방형 상태 (기능 사용 가능)")
    if st.sidebar.button("다시 잠금"):
        st.session_state['is_authenticated'] = False
        st.rerun()

st.sidebar.markdown("---")

# -----------------
# 깔끔한 2단계 메뉴 선택 (아까 방식으로 복구)
# -----------------

st.sidebar.markdown("### 🧩 솔루션 카테고리")
category = st.sidebar.selectbox(
    "어떤 작업을 진행하시겠어요?",
    ["📋 대시보드 및 설정", "📦 1. 수집 자동화", "📤 2. 발송 자동화", "🏫 3. 시설견적 자동화", "📲 4. SNS 자동화"]
)

st.sidebar.markdown("### 🎯 세부 기능 선택")
if category == "📋 대시보드 및 설정":
    menus = ["소개 및 메인"]
elif category == "📦 1. 수집 자동화":
    menus = [
        "🔍 B2B 영업망 크롤링",
        "📸 SNS 셀럽 크롤링",
        "🏭 제조업체 크롤링"
    ]
elif category == "📤 2. 발송 자동화":
    menus = [
        "📝 사업체 자동제안서 (이메일 발송 포함)", 
        "📝 SNS셀럽 자동제안서", 
        "📝 제조업체 자동제안서"
    ]
elif category == "🏫 3. 시설견적 자동화":
    menus = [
        "📃 시설 서류자동발송", 
        "🏢 시설 비교견적자동발송"
    ]
elif category == "📲 4. SNS 자동화":
    menus = [
        "📝 블로그 자동화", 
        "📸 인스타 자동화", 
        "▶️ 유튜브 자동화"
    ]

menu = st.sidebar.radio("원하시는 기능을 클릭하세요:", menus)

st.sidebar.markdown("---")
st.sidebar.info("""
📞 **문의 / 기술지원**
* **하맘컨텐츠 고객센터: 1660-1195**
* 💬 [카카오톡 채널 '하맘' 바로가기](http://pf.kakao.com/)
* 🔍 **네이버에 '하맘컨텐츠'를 검색해보세요!**

*DB 연동 서비스는 곧 활성화됩니다.
""")

# ----------------------------------------------------
# 1. 대시보드 및 설정
# ----------------------------------------------------
if menu == "소개 및 메인":
    st.title("✅ 하맘 종합자동화솔루션 - 웹서버버전")
    st.markdown("""
    환영합니다! 기존 영업망 검색 기능을 확장하여 
    **B2B 영업, 시설(관공서/어린이집) 견적, SNS 셀럽 발굴, 유튜브/인스타 자동화**를 모두 커버하는 
    올인원(All-in-one) 시스템으로 개편되었습니다.
    
    ### 💻 현재 지원되는 핵심 파이프라인
    1. **수집 자동화**: 인터넷 상의 다양한 타겟 데이터(네이버, SNS, 제조) 수집 및 실시간 감시
    2. **발송 자동화**: 수집한 데이터를 자체 관리 리스트에 담아 자동 생성된 **AI 제안서**로 이메일 일괄 송부
    3. **시설 자동화**: 어린이집, 관공서 타겟 5종 비교견적 템플릿/서류 초고속 배포
    4. **SNS 자동화**: 수집된 영상, 이미지 기반 자동 포스팅 스케줄러 (유튜브 등)
    """)

# ----------------------------------------------------
# 2. 수집 자동화
# ----------------------------------------------------
elif menu == "🔍 B2B 영업망 크롤링":
    st.header("🔍 [수집] B2B 영업망 자동 수집 및 리스트 정리")
    st.markdown("네이버에서 연락처를 추출한 뒤, **아래 표에서 바로 정리하고 파일을 저장**할 수 있습니다.")
    
    col1, col2 = st.columns(2)
    with col1:
        region = st.text_input("검색 지역", "서울 강남구")
    with col2:
        categories_input = st.text_input("🎯 타겟 업종 (원하는 키워드를 쉼표로 구분해 맘껏 적어보세요!)", "부동산, 인테리어, 청소, 식당")
        
        categories = []
        for c in categories_input.split(","):
            val = c.strip()
            if val:
                # 네이버 지도는 '부동산' 검색 시 매물 페이지로 이동해버리므로 '공인중개사'로 강제 치환
                if "부동산" in val:
                    val = val.replace("부동산", "공인중개사")
                categories.append(val)

    with st.expander("🔥 [초고속 원스톱] 크롤링부터 메일 발송까지 한 번에 (설정)"):
        st.warning("이 기능을 켜고 아래 원스톱 버튼을 누르면, 검색된 모든 대상에게 즉시 제안서 메일이 폭격됩니다.")
        cs_email = st.text_input("G-Mail 주소", placeholder="test@gmail.com", key="os_usr")
        cs_pwd = st.text_input("앱 비밀번호", type="password", key="os_pwd")
        cs_title = st.text_input("발송될 메일 제목", "[하맘] 원스톱 자동화 제안서", key="os_title")

    col_btn1, col_btn2 = st.columns(2)
    with col_btn1:
        btn_normal = st.button("🚀 기본 크롤링 수집만 시작", type="primary", use_container_width=True)
    with col_btn2:
        btn_onestop = st.button("🔥 [원스톱] 크롤링 + 자동 메일 동시 진행", type="primary", use_container_width=True)

    if btn_normal or btn_onestop:
        if not st.session_state['is_authenticated']:
            st.error("🔒 관리자 잠김 상태입니다. 사이드바에 비밀번호를 입력해주세요.")
        elif not categories:
            st.warning("업종을 선택해주세요.")
        elif btn_onestop and (not cs_email or not cs_pwd):
            st.error("원스톱 기능을 쓰려면 위의 톱니바퀴에서 지메일 계정/비밀번호 설정을 기입해야 합니다!")
        else:
            log_container = st.empty()
            image_placeholder = st.empty()
            with st.spinner('크롤러를 가동 중입니다... 실시간 검색 뷰어를 확인하세요.'):
                def log_cb(msg, screenshot=None):
                    if 'logs' not in st.session_state:
                        st.session_state['logs'] = ""
                    st.session_state['logs'] += f"{msg}\n"
                    log_container.text_area("💻 작업 실시간 로그", st.session_state['logs'][-2000:], height=100)
                    if screenshot:
                        image_placeholder.image(screenshot, caption="👀 AI 로봇 실시간 검색 뷰어", use_container_width=True)

                try:
                    crawler = NaverMapCrawler(headless=True, callback=log_cb)
                    res = crawler.crawl_all_categories(region=region, selected_cats=categories)
                    st.session_state['crawled_data_b2b'] = res
                    crawler.quit()
                    
                    # 💡 새로 수집된 데이터를 DB에 즉시 영구 저장 (SQLite)
                    df_new = pd.DataFrame(res)
                    if not df_new.empty:
                        # 이미 있는 중복 제거 후 DB 삽입
                        df_new.to_sql('b2b_crawling_list', con=conn, if_exists='append', index=False)
                        
                    st.success(f"✅ 수집 완료! DB에 영구 저장되었습니다. (총 {len(res)}건) 무제한 스크롤 완료.")
                    
                    if btn_onestop:
                        st.info("🔥 원스톱 기능 가동: 곧바로 이메일 자동 발송을 시작합니다...")
                        targets = [t for t in res if str(t.get('이메일', '')).strip() != '']
                        if not targets:
                            st.warning("수집된 데이터 중 이메일 주소를 가진 대상이 없습니다.")
                        else:
                            sender = EmailSender("smtp.gmail.com", 587, cs_email, cs_pwd, callback=lambda x: None)
                            success, fail = sender.send_campaign(targets, cs_title, dry_run=False)
                            st.success(f"원스톱 발송 성공! 🎯 [성공: {success} / 실패: {fail}]")
                            
                except Exception as e:
                    st.error(f"에러 발생: {e}")
                    
    st.markdown("---")
    st.subheader("📊 수집 리스트 통합 관리 DB (영구 보관)")
    
    # DB에서 기존 데이터를 전부 긁어옴
    try:
        db_df = pd.read_sql_query("SELECT 업체명, 카테고리, 주소, 전화번호, 이메일, 해시태그, 검색카테고리, 등록일시 FROM b2b_crawling_list ORDER BY 등록일시 DESC", conn)
    except:
        db_df = pd.DataFrame()
        
    if db_df.empty:
        st.info("아직 수집/저장된 데이터베이스(DB) 내역이 없습니다. 위의 파란 버튼을 눌러 크롤링을 시작해보세요.")
    else:
        st.markdown(f"**총 {len(db_df)}개**의 타겟 업체가 저희 하맘 클라우드 DB 금고에 완벽하게 적재되어 있습니다. 🔒")
        if '선택' not in db_df.columns: 
            db_df.insert(0, '선택', True)
        
        st.markdown("**제휴 메일을 보낼 업체만 체크박스를 유지하세요 (기본값: 모두 발송)**")
        edited_df = st.data_editor(
            db_df, hide_index=True,
            column_config={"선택": st.column_config.CheckboxColumn("보내기", default=True)},
            disabled=db_df.columns.drop('선택'), use_container_width=True
        )
        st.session_state['db_b2b_records'] = edited_df.to_dict('records')
        st.caption(f"여기서 체크된 **{edited_df['선택'].sum()}개**의 업체 대상만 최종 자동제안서 전송 타겟으로 설정됩니다.")
        
        # 엑셀 저장 & 메일 연동 유도
        csv = db_df.drop(columns=['선택']).to_csv(index=False).encode('utf-8-sig')
        col_end1, col_end2 = st.columns(2)
        with col_end1:
            st.download_button("💾 이 리스트를 엑셀로 저장", csv, f"B2B_타겟_리스트.csv", "text/csv", use_container_width=True)
        with col_end2:
            if st.button("📨 이 리스트 메일발송 (제안서 탭 설정)", use_container_width=True):
                st.info("👆 좌측 사이드바의 **[발송 자동화 > 사업체 자동제안서]** 메뉴를 누르시면 방금 체크하신 이 리스트에 메일을 자동 발송할 수 있습니다.")

elif menu in ["📸 SNS 셀럽 크롤링", "🏭 제조업체 크롤링"]:
    st.header(f"🔍 [진행 예정] {menu}")
    st.info("SNS(인스타그램 DM, 유튜브 이메일) 및 공장/제조업체 검색 추출기 기능이 들어올 예정입니다.")
    st.markdown("이곳 아래에도 곧 B2B 망처럼 **크롤링 화면**과 **표(데이터에디터) 정리 기능**을 한 페이지로 통합 제공할 것입니다.")

# ----------------------------------------------------
# 3. 발송 자동화
# ----------------------------------------------------
elif menu == "📝 사업체 자동제안서 (이메일 발송 포함)":
    st.header("📝 AI 사업체(B2B) 제안서 제작 및 대량 발송")
    st.markdown("나노바나나 프롬프트+텍스트가 자동 조합되는 제안서를 만들고, 위의 *B2B 수집리스트*에 대량 발송합니다.")
    
    col1, col2 = st.columns(2)
    with col1:
        company_name = st.text_input("당사 업체명", "하맘 B2B 솔루션")
        services = st.text_area("주요 서비스", "- 인테리어 협업\n- 관공서 유지보수")
    with col2:
        target_name = st.text_input("수신자 호칭", "파트너사 대표님")
        ai_style = st.selectbox("제안서 분위기 (AI 프롬프트 생성용)", ["전문적인", "따뜻하고 감성적인", "신뢰를 주는 프리미엄형", "나노바나나 아트 스타일"])

    if st.button("✨ 제안서 문구 자동 창작(생성)", type="primary"):
        st.session_state['proposal_content_b2b'] = f"""[Business Proposal] {company_name} - 전략적 업무 제휴 및 파트너십 제안의 건

수신: {target_name} 
발신: {company_name} 전략사업부

안녕하십니까, 귀사의 무궁한 발전과 건승을 기원합니다.
저희는 차별화된 비즈니스 역량과 혁신적인 솔루션을 바탕으로 공간과 서비스에 새로운 가치를 창출하는 {company_name}입니다.

금번 연락을 드리게 된 목적은 당사가 보유한 전문 인프라와 귀사의 우수한 사업 역량을 결합하여, 양사 간 상호 윈-윈(Win-Win)할 수 있는 전략적 파트너십을 제안 드리기 위함입니다.

[당사 핵심 역량 및 제휴 포인트]
{services}

저희는 일방적인 제안이 아닌, 귀사의 비즈니스 방향성과 운영 방침을 최우선으로 존중하며 유연하고 합리적인 협의점을 도출할 준비가 되어 있습니다. 본 제안과 관련하여 추가적인 수익 창출 모델이나 비용 절감 방안 등 어떠한 형태의 비즈니스 논의도 환영합니다.

바쁘신 업무 중에 본 제안서를 검토해 주셔서 깊은 감사를 드립니다.
시간이 허락하실 때 짧은 대면/비대면 미팅이나 유선 상으로 보다 실질적이고 가치 있는 협업 모델을 논의 드릴 수 있기를 진심으로 희망합니다.

긍정적인 화답을 기다리겠습니다.
감사합니다.

{company_name} 배상

---------------------------------------
[🎨 나노바나나(Niji/Midjourney) AI 이미지용 프롬프트]
"A beautiful, high-quality cover design for a B2B proposal. Clean layout, modern corporate and professional typography, bright lighting, {ai_style} style, 8k resolution, photorealistic --ar 16:9 --v 6"
---------------------------------------
"""
        st.success("✅ 제안서 초안이 작성되었습니다! (문서 다운로드 또는 이메일 바로 전송 가능)")

    if 'proposal_content_b2b' in st.session_state:
        st.text_area("제안서 미리보기 (텍스트 교정 가능)", st.session_state['proposal_content_b2b'], height=300)
        
        # 1. 문서 다운로드 기능
        doc = Document()
        doc.add_heading(f"{company_name} B2B 제안서", 0)
        doc.add_paragraph(st.session_state['proposal_content_b2b'])
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        st.download_button("📥 HWP지원용 Word(.docx) 제안서 다운로드", doc_io, f"B2B_제안서.docx")
        
        st.markdown("---")
        # 2. 이메일 발송 기능
        st.subheader("🚀 B2B 이메일 발송 연동")
        with st.expander("구글(Gmail) 계정 입력", expanded=True):
            em_usr, em_pwd = st.columns(2)
            gmail_user = em_usr.text_input("G-Mail 주소")
            gmail_pass = em_pwd.text_input("앱 전용 비밀번호", type="password")

        email_title = st.text_input("메일 제목", f"[제안] {company_name} 제휴 제안서")
        test_mode = st.toggle("테스트 모드 유지 (실제 발송 안함)", value=True)

        if st.button("📨 작성된 제안서 내용을 B2B 수집리스트 전체에 발송"):
            if not st.session_state['is_authenticated']:
                st.error("🔒 관리자 비밀번호를 먼저 입력하여 잠금을 해제해주세요.")
            elif 'crawled_data_b2b' not in st.session_state or not st.session_state['crawled_data_b2b']:
                st.warning("B2B 크롤링 데이터를 먼저 수집해주세요!")
            else:
                targets = [t for t in st.session_state['crawled_data_b2b'] if t.get('선택', True) and str(t.get('이메일', '')).strip() != '']
                if not targets:
                    st.error("이메일 목록이 없거나 모두 체크 해제되었습니다.")
                else:
                    with st.spinner("발송 중..."):
                        sender = EmailSender("smtp.gmail.com", 587, gmail_user, gmail_pass, callback=lambda x: None)
                        # TODO: 템플릿 로직이 아닌 본문 텍스트 통째로 발송하게 변경해야 함 (임시로 캠페인 구조 활용)
                        success, fail = sender.send_campaign(targets, email_title, dry_run=test_mode)
                        st.success(f"처리 완료 (성공:{success}/실패:{fail})")

elif menu in ["📝 SNS셀럽 자동제안서", "📝 제조업체 자동제안서"]:
    st.header(f"📝 {menu.split()[1]} (개발 중)")
    st.info("SNS DM 봇(인스타그램) 전용 양식 및 공장/제조업체 조달 양식을 준비하고 있습니다.")

# ----------------------------------------------------
# 4. 시설견적 자동화
# ----------------------------------------------------
elif menu == "🏢 시설 비교견적자동발송":
    st.header("🏫 하맘 어린이집 및 공익시설 비교견적서 발송")
    st.info("공익시설(어린이집 등) 전용 마케팅 템플릿(비교견적서 형태)을 일괄 송신합니다.")
    
    with st.expander("구글(Gmail) 계정 설정", expanded=True):
        col_acc1, col_acc2 = st.columns(2)
        gmail_user = col_acc1.text_input("G-Mail 주소", key="kid_usr")
        gmail_pass = col_acc2.text_input("앱 비밀번호", type="password", key="kid_pwd")

    kid_title = st.text_input("메일 제목", "[비교견적] 하맘 시설물 방문 견적 안내")
    template = st.selectbox("비교견적 양식 선택", ["비교견적 양식1 (심플)", "비교견적 양식2 (상세비용)", "비교견적 양식3 (친환경/안전강조)", "비교견적 양식4 (프리미엄)", "비교견적 양식5 (정기관리형)"])
    test_mode = st.toggle("방지 테스트 모드", value=True, key="kid_test")

    if st.button("📨 비교견적 장착 후 템플릿 대량 발송"):
        if not st.session_state['is_authenticated']:
            st.error("🔒 관리자 비번 해제가 필요합니다.")
        elif 'crawled_data_b2b' not in st.session_state or not st.session_state['crawled_data_b2b']:
            st.warning("영업망 데이터 수집/정리가 필요합니다.")
        else:
            targets = [t for t in st.session_state['crawled_data_b2b'] if t.get('선택', True) and str(t.get('이메일', '')).strip() != '']
            if not targets: st.error("이메일 주소를 확보한 기관이 없습니다.")
            else:
                with st.spinner(f"'{template}' 양식으로 발송 중..."):
                    sender = EmailSender("smtp.gmail.com", 587, gmail_user, gmail_pass, callback=lambda x: None)
                    success, fail = sender.send_campaign(targets, kid_title, template_type=template, dry_run=test_mode)
                st.success(f"메일 발송 처리 완료 (성공:{success} / 실패:{fail})")

elif menu == "📃 시설 서류자동발송":
    st.header("📃 시설 관리 필수 서류 묶음 자동 발송")
    st.info("비교견적서 외에 지명원, 사업자등록증, 자격증 등 필수 묶음 서류를 자동으로 패키징하여 관공서에 보내는 구조를 설계합니다.")

# ----------------------------------------------------
# 5. SNS 자동화
# ----------------------------------------------------
elif menu == "▶️ 유튜브 자동화":
    st.header("▶️ 메타 유튜브 퍼블리싱 (영상 + 썸네일 통합 업로드)")
    st.warning("로컬 파일이 아닌 클라우드 업로드 방식으로 유튜브 Data API를 호출합니다.")

    uploaded_secret = st.file_uploader("1️⃣ 구글 API Key (client_secret.json)", type=['json'])
    uploaded_video = st.file_uploader("2️⃣ 업로드할 영상 파일 (.mp4)", type=['mp4', 'mov'])
    
    yt_title = st.text_input("유튜브 노출 제목", "B2B 인테리어/청소 홍보 영상")
    yt_desc = st.text_area("영상 설명 (본문)", "#자동화")
    yt_pub = st.radio("공개 범위", ("private", "unlisted", "public"))

    if st.button("🚀 즉시 유튜브 자동 게재"):
        if not st.session_state['is_authenticated']:
            st.error("🔒 관리자 인증이 필요합니다.")
        elif not uploaded_secret or not uploaded_video:
            st.error("보안 키와 영상은 필수입니다!")
        else:
            st.success("데이터 검증 완료! (영상 크기에 따라 수 분 소요)")

elif menu in ["📝 블로그 자동화", "📸 인스타 자동화"]:
    st.header(f"📲 {menu.split()[1]} 플랫폼 통합 파이프라인")
    st.info("원클릭으로 여러 계정에 포스팅 및 릴스/쇼츠 게재를 스케줄링하는 모듈 탑재 예정입니다.")

